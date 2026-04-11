"""文档导入管道：批量清洗并导入 .docx 文件到知识库。

提供两个核心类：
  - DocCleaner: 使用 LLM 辅助检测，清洗 .docx 文件中的非电商内容
  - BulkImporter: 批量处理目录中的 .docx 文件，清洗→分块→嵌入→入库

典型用法::

    cleaner = DocCleaner()
    importer = BulkImporter(cleaner)
    result = await importer.import_directory(
        r"F:\\跨境电商长期主义",
        sample_mode=True,
        sample_size=20,
    )
    print(result)
"""

from __future__ import annotations

import asyncio
import random
import re
import uuid
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable

try:
    from loguru import logger  # pyright: ignore[reportMissingImports]
except ImportError:
    import logging as _logging

    logger = _logging.getLogger(__name__)  # type: ignore[assignment]

# ---------------------------------------------------------------------------
# 可选依赖
# ---------------------------------------------------------------------------
_DocxDocument: Any = None  # 占位，导入成功后替换

try:
    from docx import Document as _DocxDocumentImpl  # type: ignore[import-untyped]

    _DocxDocument = _DocxDocumentImpl
except ImportError:
    logger.warning("python-docx 未安装，.docx 文件将无法解析")  # type: ignore[union-attr]

# ---------------------------------------------------------------------------
# 内部模块依赖（延迟导入，避免循环 / 缺少 DB 时报错）
# ---------------------------------------------------------------------------
_llm_chat: Callable[..., dict[str, Any]] | None = None

try:
    from src.llm.client import chat as _chat_impl

    _llm_chat = _chat_impl
except Exception:  # pragma: no cover
    pass

_get_engine: Callable[..., Any] | None = None

try:
    from src.knowledge_base.rag_engine import _get_engine as _engine_getter  # pyright: ignore[reportPrivateUsage]

    _get_engine = _engine_getter
except Exception:  # pragma: no cover
    pass

# ---------------------------------------------------------------------------
# 常量
# ---------------------------------------------------------------------------
_SKIP_EXTENSIONS: set[str] = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".mp4", ".avi", ".mov", ".mp3"}

_CLEAN_MODEL = "gpt-4o-mini"

_SYSTEM_PROMPT_CLEAN = """\
你是一个专业的跨境电商文档清洗助手。你的任务是从原始文本中：

1. **删除** 以下无关内容：
   - 个人感悟、人生哲理、心灵鸡汤、生活随笔
   - 二维码引用、扫码关注、关注公众号等推广内容
   - 作者署名、微信号、公众号名称、个人简介
   - 广告、课程推销、付费社群推广
   - 与电商/亚马逊运营完全无关的段落

2. **保留** 以下有价值内容：
   - 亚马逊运营知识（选品、Listing、广告、物流等）
   - 跨境电商策略、数据分析、市场调研
   - 操作教程、规则解读、案例分析
   - 文档的原始结构（标题层级、列表、表格描述）

3. **输出规则**：
   - 如果清洗后仍有有价值的电商内容，直接输出清洗后的文本（不要加任何额外解释）
   - 如果整篇文档完全不含电商相关内容，只输出一个词: EMPTY
   - 保持原文的中文或英文语言
   - 不要添加你自己的评论或总结
"""

# 分块参数
_CHUNK_MAX_CHARS = 3200  # ≈ 800 tokens
_CHUNK_OVERLAP_CHARS = 400  # ≈ 100 tokens


# ---------------------------------------------------------------------------
# ImportResult 数据类
# ---------------------------------------------------------------------------
@dataclass(slots=True)
class ImportResult:
    """批量导入的结果统计。"""

    processed: int = 0
    skipped: int = 0
    failed: int = 0
    total: int = 0
    sample_files: list[str] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)

    def __str__(self) -> str:
        return (
            f"ImportResult(total={self.total}, processed={self.processed}, "
            f"skipped={self.skipped}, failed={self.failed}, "
            f"errors={len(self.errors)})"
        )


# ---------------------------------------------------------------------------
# DocCleaner — LLM 辅助文档清洗
# ---------------------------------------------------------------------------
class DocCleaner:
    """使用 LLM 辅助检测清洗 .docx 文件，去除非电商内容。

    清洗步骤（全部由 LLM 一次性完成）：
      1. 移除个人感悟 / 生活随笔
      2. 移除二维码引用与推广内容
      3. 移除作者署名与微信 / 公众号引用
      4. 仅提取电商 / 亚马逊相关内容
      5. 保留文档结构（标题、列表、表格）
    """

    def __init__(self, model: str = _CLEAN_MODEL) -> None:
        self._model = model

    # ------------------------------------------------------------------
    # 公共接口
    # ------------------------------------------------------------------
    async def clean(self, docx_path: str) -> str | None:
        """清洗单个 .docx 文件，返回清洗后文本；完全无关时返回 ``None``。

        Args:
            docx_path: .docx 文件的绝对或相对路径。

        Returns:
            清洗后的纯文本，如果整篇文档不含电商内容则返回 ``None``。

        Raises:
            ImportError: python-docx 未安装。
            FileNotFoundError: 文件不存在。
        """
        raw_text = self._extract_text(docx_path)
        if not raw_text or not raw_text.strip():
            logger.info("文件内容为空，跳过: %s", docx_path)
            return None

        cleaned = await self._llm_clean(raw_text)
        return cleaned

    # ------------------------------------------------------------------
    # 内部：读取 .docx
    # ------------------------------------------------------------------
    @staticmethod
    def _extract_text(docx_path: str) -> str:
        """从 .docx 文件中提取纯文本（保留段落结构）。"""
        if _DocxDocument is None:
            raise ImportError("python-docx 未安装，无法解析 .docx 文件")

        path = Path(docx_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {docx_path}")

        doc = _DocxDocument(str(path))

        parts: list[str] = []

        # 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # 保留标题层级标记
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading", "").strip() or "1"
                parts.append(f"{'#' * int(level)} {text}")
            else:
                parts.append(text)

        # 提取表格内容
        for table in doc.tables:
            table_rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(cells))
            if table_rows:
                parts.append("\n".join(table_rows))

        return "\n\n".join(parts)

    # ------------------------------------------------------------------
    # 内部：LLM 清洗
    # ------------------------------------------------------------------
    async def _llm_clean(self, raw_text: str) -> str | None:
        """调用 LLM 清洗原始文本。"""
        if _llm_chat is None:
            raise ImportError("src.llm.client.chat 不可用，无法执行 LLM 清洗")

        chat_fn = _llm_chat  # 局部绑定，帮助类型推断

        # 截断超长文本以控制成本（保留前 12000 字符 ≈ 3000 tokens 输入）
        truncated = raw_text[:12000]

        messages: list[dict[str, str]] = [
            {"role": "system", "content": _SYSTEM_PROMPT_CLEAN},
            {"role": "user", "content": truncated},
        ]

        # chat() 是同步函数，在线程池中运行以避免阻塞事件循环
        loop = asyncio.get_running_loop()
        response = await loop.run_in_executor(
            None,
            lambda: chat_fn(
                model=self._model,
                messages=messages,
                temperature=0.1,
                max_tokens=4000,
            ),
        )

        content: str = response.get("content", "").strip()

        if not content or content.upper() == "EMPTY":
            return None

        # 基础后处理：压缩多余空行
        content = re.sub(r"\n{3,}", "\n\n", content)
        return content


# ---------------------------------------------------------------------------
# BulkImporter — 批量导入
# ---------------------------------------------------------------------------
class BulkImporter:
    """批量导入 .docx 文件到知识库。

    流程：扫描目录 → 过滤非 .docx → DocCleaner 清洗 → 分块 → 嵌入 → 入库。
    支持 ``sample_mode`` 仅处理少量文件供老板预览。
    """

    def __init__(self, cleaner: DocCleaner | None = None) -> None:
        self._cleaner = cleaner or DocCleaner()

    # ------------------------------------------------------------------
    # 公共接口
    # ------------------------------------------------------------------
    async def import_directory(
        self,
        directory: str,
        sample_mode: bool = False,
        sample_size: int = 20,
    ) -> ImportResult:
        """批量导入目录下的 .docx 文件。

        Args:
            directory: 包含 .docx 文件的目录路径。
            sample_mode: 为 ``True`` 时仅随机抽取 ``sample_size`` 个文件处理。
            sample_size: 抽样数量（仅在 ``sample_mode=True`` 时生效）。

        Returns:
            :class:`ImportResult` 统计对象。
        """
        dir_path = Path(directory)
        if not dir_path.exists() or not dir_path.is_dir():
            raise FileNotFoundError(f"目录不存在: {directory}")

        # 收集所有文件
        all_files = list(dir_path.iterdir())
        docx_files: list[Path] = []
        skipped_count = 0

        for f in all_files:
            if not f.is_file():
                continue
            ext = f.suffix.lower()
            if ext == ".docx":
                docx_files.append(f)
            elif ext in _SKIP_EXTENSIONS or ext != ".docx":
                skipped_count += 1

        # 抽样模式
        if sample_mode and len(docx_files) > sample_size:
            docx_files = random.sample(docx_files, sample_size)
            logger.info(
                "抽样模式：从 %d 个 .docx 文件中随机选取 %d 个",
                len(list(dir_path.glob("*.docx"))),
                sample_size,
            )

        result = ImportResult(
            total=len(docx_files) + skipped_count,
            skipped=skipped_count,
        )

        logger.info(
            "开始导入: 目录=%s, .docx文件=%d, 跳过=%d, sample_mode=%s",
            directory,
            len(docx_files),
            skipped_count,
            sample_mode,
        )

        for i, docx_path in enumerate(docx_files, 1):
            filename = docx_path.name
            try:
                logger.info("[%d/%d] 处理: %s", i, len(docx_files), filename)

                # 1. 清洗
                cleaned_text = await self._cleaner.clean(str(docx_path))
                if cleaned_text is None:
                    logger.info("文件不含电商内容，跳过: %s", filename)
                    result.skipped += 1
                    continue

                # 2. 分块
                chunks = self._split_chunks(cleaned_text, source=str(docx_path))

                # 3. 嵌入 + 入库
                await self._embed_and_store(
                    chunks=chunks,
                    source=str(docx_path),
                    filename=filename,
                )

                result.processed += 1
                if sample_mode:
                    result.sample_files.append(filename)

                logger.info(
                    "导入成功: %s (%d 块) [%d/%d]",
                    filename,
                    len(chunks),
                    result.processed,
                    len(docx_files),
                )

            except Exception as exc:
                error_msg = f"{filename}: {exc}"
                logger.error("导入失败: %s", error_msg)
                result.failed += 1
                result.errors.append(error_msg)

        logger.info(
            "导入完成: %s",
            result,
        )
        return result

    # ------------------------------------------------------------------
    # 内部：文本分块
    # ------------------------------------------------------------------
    @staticmethod
    def _split_chunks(text: str, source: str) -> list[dict[str, object]]:
        """将清洗后的文本按段落分块。

        Args:
            text: 清洗后的文本。
            source: 原始文件路径（用于元数据）。

        Returns:
            分块列表，每块包含 ``chunk_text``, ``chunk_index``, ``metadata``。
        """
        paragraphs = [p.strip() for p in re.split(r"\n\n+", text) if p.strip()]

        chunks: list[dict[str, object]] = []
        current_parts: list[str] = []
        current_len = 0
        chunk_idx = 0

        now_str = datetime.now(tz=timezone.utc).isoformat()

        def _flush() -> None:
            nonlocal chunk_idx
            joined = "\n\n".join(current_parts).strip()
            if joined:
                chunks.append(
                    {
                        "chunk_text": joined,
                        "chunk_index": chunk_idx,
                        "metadata": {
                            "source": source,
                            "title": Path(source).stem,
                            "category": "导入文档",
                            "doc_type": "other",
                            "import_date": now_str,
                            "original_filename": Path(source).name,
                        },
                    }
                )
                chunk_idx += 1

        for para in paragraphs:
            para_len = len(para)
            if current_len + para_len > _CHUNK_MAX_CHARS and current_parts:
                _flush()
                # 保留重叠
                overlap_text = "\n\n".join(current_parts)[-_CHUNK_OVERLAP_CHARS:]
                current_parts = [overlap_text] if overlap_text else []
                current_len = len(overlap_text)

            current_parts.append(para)
            current_len += para_len

        # 最后一块
        if current_parts:
            _flush()

        # 空文本保底
        if not chunks:
            chunks.append(
                {
                    "chunk_text": text,
                    "chunk_index": 0,
                    "metadata": {
                        "source": source,
                        "title": Path(source).stem,
                        "category": "导入文档",
                        "doc_type": "other",
                        "import_date": now_str,
                        "original_filename": Path(source).name,
                    },
                }
            )

        return chunks

    # ------------------------------------------------------------------
    # 内部：嵌入并存储
    # ------------------------------------------------------------------
    async def _embed_and_store(
        self,
        chunks: list[dict[str, object]],
        source: str,
        filename: str,
    ) -> None:
        """生成向量嵌入并写入知识库。"""
        if _get_engine is None:
            raise ImportError("rag_engine 不可用，无法嵌入和存储文档")

        engine = _get_engine()

        # 生成确定性 doc_id
        doc_id = uuid.uuid5(uuid.NAMESPACE_URL, source)

        loop = asyncio.get_running_loop()

        # ingest_chunks 是同步方法 — 在线程池中运行
        await loop.run_in_executor(
            None,
            lambda: engine.ingest_chunks(chunks),
        )

        logger.debug(
            "已存储 %d 块到知识库: %s (doc_id=%s)",
            len(chunks),
            filename,
            doc_id,
        )
